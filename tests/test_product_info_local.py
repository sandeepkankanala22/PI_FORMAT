"""Local verification tests for Product Information auto-fill."""

from __future__ import annotations

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from fastapi.testclient import TestClient

from backend.server import app
from backend.services.product_info.document_text_extraction import DocumentType, detect_document_type, extract_text_from_bytes
from backend.services.product_info.field_mapping import map_to_form_fields
from backend.services.product_info.field_summarization import needs_indication_refinement
from backend.services.product_info.json_validation import parse_json_from_response, validate_product_info_json
from backend.services.product_info.text_chunking import chunk_text_if_needed
from backend.services.product_info.text_cleaning import clean_text, prepare_pi_text_for_llm

client = TestClient(app)


def test_service_modules() -> None:
    raw = "Page 1\n\nProduct: TUB-040\nIndication: NSCLC\n\nPage 2\n\nProduct: TUB-040\n"
    assert "TUB-040" in clean_text(raw)

    noisy = (
        "Product Name: TUB-040\n"
        "Indication: NSCLC\n"
        "https://example.com/label.pdf\n"
        "REFERENCES\n"
        "Smith et al. 2020 https://doi.org/10.0000/example\n"
        "ADVERSE REACTIONS\n"
        "nausea vomiting fatigue\n"
        "Mechanism of Action: TOP1 inhibitor\n"
    )
    prepared = prepare_pi_text_for_llm(noisy)
    assert "TUB-040" in prepared
    assert "NSCLC" in prepared
    assert "TOP1 inhibitor" in prepared
    assert "https://" not in prepared
    assert len(prepared) < len(noisy)

    long_text = "indication " * 5000
    assert len(chunk_text_if_needed(long_text, char_limit=1000)) <= 1000

    llm = (
        '{"product_name": "TUB-040", "country": "USA", "indication": "NSCLC", '
        '"drug_class": "ADC", "mechanism_of_action": "TOP1 inhibitor", '
        '"start_year": "2025", "end_year": "2035"}'
    )
    model = validate_product_info_json(parse_json_from_response(llm))
    fields = map_to_form_fields(model)
    assert fields["productName"] == "TUB-040"
    assert fields["country"] == "United States"
    assert "ADC" in fields["classMoa"]
    assert "launchYear" not in fields
    assert "peakYear" not in fields

    doc = Document()
    doc.add_paragraph("Product Name: TestDrug")
    doc.add_paragraph("Indication: Rheumatoid Arthritis")
    buf = io.BytesIO()
    doc.save(buf)
    assert "TestDrug" in extract_text_from_bytes(buf.getvalue(), "test.docx")
    assert detect_document_type(b"%PDF-1.4", "a.pdf") == DocumentType.PDF

    statin_label = (
        "Prevention of cardiovascular disease (reduce risk of MI, stroke, revascularization procedures, angina); "
        "treatment of hyperlipidemia, mixed dyslipidemia, hypertriglyceridemia, primary dysbetalipoproteinemia, "
        "homozygous familial hypercholesterolemia, heterozygous familial hypercholesterolemia"
    )
    assert needs_indication_refinement(statin_label)
    assert not needs_indication_refinement("Hyperlipidemia")
    assert not needs_indication_refinement("Non-small cell lung cancer")


def test_api_endpoints() -> None:
    assert client.post("/api/product-info/extract").status_code == 400

    r = client.post(
        "/api/product-info/extract",
        files={"file": ("bad.txt", b"hello", "text/plain")},
    )
    assert r.status_code == 400

    r = client.post(
        "/api/product-info/extract",
        files={"file": ("legacy.doc", b"fake", "application/msword")},
    )
    assert r.status_code == 400
    assert "docx" in r.json()["detail"].lower()

    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    r = client.post(
        "/api/product-info/extract",
        files={"file": ("empty.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 422

    r = client.post(
        "/api/prompts/upload-resource",
        files={"file": ("test.txt", b"hello world", "text/plain")},
    )
    assert r.status_code == 200
    assert "content" in r.json()

    r = client.get("/health")
    assert r.status_code == 200

    r = client.post(
        "/api/product-info/extract",
        json={"url": "not-a-url"},
    )
    assert r.status_code == 400


def test_live_docx_extract() -> None:
    """End-to-end Bedrock extraction with a synthetic PI-like DOCX."""
    doc = Document()
    doc.add_paragraph("PRESCRIBING INFORMATION")
    doc.add_paragraph("Product Name: OncoVax-101")
    doc.add_paragraph("Drug Class: Monoclonal Antibody")
    doc.add_paragraph("Mechanism of Action: PD-1 inhibitor")
    doc.add_paragraph("Indication: Non-small cell lung cancer (NSCLC)")
    doc.add_paragraph("Launch Year: 2026")
    doc.add_paragraph("Forecast End Year: 2035")
    doc.add_paragraph("Country: United States")
    buf = io.BytesIO()
    doc.save(buf)

    r = client.post(
        "/api/product-info/extract",
        files={"file": ("pi-sample.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    print("live extract status:", r.status_code)
    data = r.json()
    if r.status_code == 200:
        print("fields:", data.get("fields"))
        assert data["status"] == "ok"
        assert data.get("fields")
        assert "launchYear" not in data.get("fields", {})
        assert "peakYear" not in data.get("fields", {})
    else:
        # Bedrock may be unavailable in CI — surface detail for manual review
        print("detail:", data.get("detail"))
        if r.status_code in (500, 504):
            print("SKIP live Bedrock assertion (service unavailable)")
            return
        raise AssertionError(f"Unexpected status {r.status_code}: {data}")


if __name__ == "__main__":
    test_service_modules()
    print("service modules: PASS")
    test_api_endpoints()
    print("api endpoints: PASS")
    test_live_docx_extract()
    print("live docx extract: PASS (or skipped if Bedrock unavailable)")
    print("\nAll local tests completed.")
