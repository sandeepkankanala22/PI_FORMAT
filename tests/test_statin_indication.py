"""Quick live test for statin-style verbose indication refinement."""
import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from fastapi.testclient import TestClient
from backend.server import app

client = TestClient(app)

doc = Document()
doc.add_paragraph("Product Name: Lipitor-Stat")
doc.add_paragraph("Drug Class: HMG-CoA reductase inhibitor")
doc.add_paragraph("Mechanism of Action: Statin")
doc.add_paragraph(
    "Indications and Usage: Prevention of cardiovascular disease (reduce risk of MI, stroke, "
    "revascularization procedures, angina); treatment of hyperlipidemia, mixed dyslipidemia, "
    "hypertriglyceridemia, primary dysbetalipoproteinemia, homozygous familial hypercholesterolemia, "
    "heterozygous familial hypercholesterolemia"
)
doc.add_paragraph("Country: United States")
doc.add_paragraph("Launch Year: 2025")
doc.add_paragraph("Forecast End Year: 2035")
buf = io.BytesIO()
doc.save(buf)

r = client.post(
    "/api/product-info/extract",
    files={"file": ("statin-pi.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
)
print("status:", r.status_code)
data = r.json()
print("indication:", data.get("fields", {}).get("indication"))
print("all fields:", data.get("fields"))
