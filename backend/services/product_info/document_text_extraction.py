"""Document type detection and text extraction (PDF via PyMuPDF, DOCX via python-docx)."""

from __future__ import annotations

import io
import logging
from enum import Enum
from typing import Optional

from .errors import CorruptDocument, EmptyDocument, UnsupportedFileType

logger = logging.getLogger("ProductInfoExtract")


class DocumentType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    UNKNOWN = "unknown"


def detect_document_type(raw_bytes: bytes, filename: str = "") -> DocumentType:
    """Detect document type from extension and magic bytes."""
    fname_lower = (filename or "").lower()
    if fname_lower.endswith(".pdf") or raw_bytes[:4] == b"%PDF":
        return DocumentType.PDF
    if fname_lower.endswith(".docx") or raw_bytes[:2] == b"PK":
        if fname_lower.endswith(".doc") and not fname_lower.endswith(".docx"):
            return DocumentType.DOC
        return DocumentType.DOCX
    if fname_lower.endswith(".doc"):
        return DocumentType.DOC
    return DocumentType.UNKNOWN


def extract_text_from_bytes(raw_bytes: bytes, filename: str = "") -> str:
    """Extract readable text from PDF or DOCX bytes."""
    doc_type = detect_document_type(raw_bytes, filename)
    logger.info("Document type detected: %s (filename=%s, size=%d)", doc_type.value, filename, len(raw_bytes))

    if doc_type == DocumentType.DOC:
        raise UnsupportedFileType("Please convert .doc to .docx and upload again.")
    if doc_type == DocumentType.UNKNOWN:
        raise UnsupportedFileType("Unsupported file type. Use PDF or DOCX.")
    if doc_type == DocumentType.PDF:
        return _extract_pdf(raw_bytes)
    if doc_type == DocumentType.DOCX:
        return _extract_docx(raw_bytes)
    raise UnsupportedFileType("Unsupported file type. Use PDF or DOCX.")


def _extract_pdf(raw_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise CorruptDocument("PDF extraction requires pymupdf. Run: pip install pymupdf") from exc

    try:
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        parts: list[str] = []
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                parts.append(text.strip())
        doc.close()
    except Exception as exc:
        logger.error("PDF extraction failed: %s", exc, exc_info=True)
        raise CorruptDocument(f"Could not read PDF: {exc}") from exc

    combined = "\n\n".join(parts).strip()
    if not combined:
        raise EmptyDocument("No readable text found in PDF. Scanned/image PDFs are not supported.")
    return combined


def _extract_docx(raw_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise CorruptDocument("DOCX extraction requires python-docx. Run: pip install python-docx") from exc

    try:
        doc = Document(io.BytesIO(raw_bytes))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc, exc_info=True)
        raise CorruptDocument(f"Could not read DOCX: {exc}") from exc

    combined = "\n\n".join(parts).strip()
    if not combined:
        raise EmptyDocument("No readable text found in DOCX.")
    return combined
